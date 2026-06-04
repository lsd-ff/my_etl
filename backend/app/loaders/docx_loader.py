"""DOCX loader using python-docx."""

from __future__ import annotations

from pathlib import Path

from app.loaders.base_loader import BaseLoader
from app.schemas.qa_record import LoadedSegment


class DocxLoader(BaseLoader):
    supported_extensions = {".docx"}

    def load(self, file_path: str | Path) -> list[LoadedSegment]:
        path = Path(file_path)
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("DOCX support requires python-docx. Install requirements.txt.") from exc

        document = Document(str(path))
        segments: list[LoadedSegment] = []
        current_lines: list[str] = []
        heading_path: list[str] = []
        section = ""
        offset = 0

        def flush(block_type: str = "text") -> None:
            nonlocal current_lines, offset
            text = "\n".join(current_lines).strip()
            if not text:
                current_lines = []
                return
            start = offset
            end = start + len(text)
            segments.append(
                LoadedSegment(
                    text=text,
                    source=str(path),
                    file_type="docx",
                    page=0,
                    section=section,
                    heading_path=tuple(heading_path),
                    block_type=block_type,
                    start_offset=start,
                    end_offset=end,
                )
            )
            offset = end + 1
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style_name or "标题" in style_name:
                flush()
                section = text
                level = self._heading_level(style_name)
                heading_path[:] = heading_path[: max(0, level - 1)]
                heading_path.append(text)
                current_lines.append(text)
                flush("heading")
                continue
            current_lines.append(text)

        flush()

        for table_index, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                text = "\n".join(rows)
                start = offset
                end = start + len(text)
                segments.append(
                    LoadedSegment(
                        text=text,
                        source=str(path),
                        file_type="docx",
                        page=0,
                        section=section,
                        heading_path=tuple(heading_path),
                        block_type="table",
                        start_offset=start,
                        end_offset=end,
                        warnings=(f"table_{table_index}",),
                    )
                )
                offset = end + 1

        return segments

    @staticmethod
    def _heading_level(style_name: str) -> int:
        for token in reversed(style_name.split()):
            if token.isdigit():
                return max(1, min(6, int(token)))
        return 1
